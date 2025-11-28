import os
import re
import pandas as pd
from collections import Counter
import matplotlib.pyplot as plt
from docx import Document
from pdf2image import convert_from_path
import pytesseract
from openpyxl import Workbook, load_workbook

# =======================
# CONFIGURATION
# =======================
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
POPPLER_PATH = r"C:\poppler\poppler-25.11.0\Library\bin"  # Update if different

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# To store last analysis results
last_text_result = None
last_csv_result = None

# =======================
# TEXT FILE ANALYSIS
# =======================
def analyze_text_file(file_path):
    global last_text_result

    if not os.path.exists(file_path):
        print("❌ File not found!")
        return
    
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        
        words = text.split()
        total_words = len(words)
        unique_words = len(set(words))
        word_freq = Counter(words).most_common(1)[0]
        emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        numbers = re.findall(r'\b\d+\b', text)

        print("\n📌 TEXT FILE ANALYSIS RESULT")
        print("-" * 30)
        print(f"📄 Total Words: {total_words}")
        print(f"🔁 Unique Words: {unique_words}")
        print(f"🔥 Most Frequent Word: '{word_freq[0]}' ({word_freq[1]})")
        print(f"📧 Emails Found: {emails if emails else 'None'}")
        print(f"🔢 Numbers Found: {numbers if numbers else 'None'}")

        last_text_result = f"""
Total Words: {total_words}
Unique Words: {unique_words}
Most Frequent Word: {word_freq[0]} ({word_freq[1]})
Emails: {emails if emails else 'None'}
Numbers: {numbers if numbers else 'None'}
"""

        chart_path = generate_word_chart(words)
        generate_report(last_text_result, chart_path)

    except Exception as e:
        print(f"❌ Error: {e}")

# =======================
# CSV FILE ANALYSIS
# =======================
def analyze_csv_file(file_path):
    global last_csv_result

    if not os.path.exists(file_path):
        print("❌ File not found!")
        return
    
    try:
        df = pd.read_csv(file_path)
        print("\n📌 CSV FILE ANALYSIS RESULT")
        print("-" * 30)
        print(df.describe())
        print(f"\n📊 Total Rows: {len(df)}, Columns: {len(df.columns)}")
        print("📋 Column Names:", list(df.columns))

        last_csv_result = df.describe()

    except Exception as e:
        print(f"❌ Error: {e}")

# =======================
# GENERATE CHART
# =======================
def generate_word_chart(words_list):
    if not words_list:
        return None

    word_freq = Counter(words_list).most_common(5)
    words = [w[0] for w in word_freq]
    counts = [w[1] for w in word_freq]

    chart_path = os.path.join(BASE_DIR, "output", "charts", "word_chart.png")
    os.makedirs(os.path.dirname(chart_path), exist_ok=True)

    plt.figure()
    plt.bar(words, counts)
    plt.title("Top 5 Words")
    plt.xlabel("Words")
    plt.ylabel("Frequency")
    plt.savefig(chart_path)
    plt.close()

    print(f"📊 Chart saved at: {chart_path}")
    return chart_path

# =======================
# REPORT GENERATION
# =======================
def generate_report(result_text, chart_path=None):
    reports_dir = os.path.join(BASE_DIR, "output", "reports")
    os.makedirs(reports_dir, exist_ok=True)

    txt_path = os.path.join(reports_dir, "analysis_report.txt")
    doc_path = os.path.join(reports_dir, "analysis_report.docx")

    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(result_text)

    doc = Document()
    doc.add_heading("File Analysis Report", 0)
    doc.add_paragraph(result_text)

    if chart_path:
        doc.add_picture(chart_path)

    doc.add_paragraph("\nGenerated using Python File Analyzer Tool")
    doc.save(doc_path)

    print(f"📁 Reports saved at:\n • {txt_path}\n • {doc_path}")

# =======================
# OCR FUNCTIONS
# =======================
def is_aadhar(text):
    match = re.search(r'\b\d{4}[\s-]?\d{4}[\s-]?\d{4}\b', text)
    return match.group(0) if match else None

def extract_dob(text):
    match = re.search(r'\b(\d{2}/\d{2}/\d{4})\b', text)
    return match.group(1) if match else "DOB Not Found"

def extract_gender(text):
    gender_keywords = ["MALE", "FEMALE", "Male", "Female", "M", "F"]
    for word in gender_keywords:
        if word.lower() in text.lower():
            return word.capitalize()
    return "Gender Not Found"

def get_name_near_aadhar(text):
    lines = text.split('\n')
    for line in lines:
        if "name" in line.lower():
            clean = re.sub(r'[^A-Za-z\s]', '', line)
            return clean.replace("Name", "").strip()
    for line in lines:
        if line.strip().isupper() and len(line.split()) >= 2:
            return line.strip()
    return "Name Not Found"

def save_to_excel(name, aadhar, dob, gender):
    excel_path = os.path.join(BASE_DIR, "aadhar_output.xlsx")

    try:
        if os.path.exists(excel_path):
            wb = load_workbook(excel_path)
            ws = wb.active
        else:
            wb = Workbook()
            ws = wb.active
            ws.append(["Name", "Aadhaar Number", "DOB", "Gender"])

        ws.append([name, aadhar, dob, gender])
        wb.save(excel_path)

        print(f"📁 Data saved to Excel: {excel_path}")
    except Exception as e:
        print(f"❌ Excel Save Error: {e}")

def extract_aadhar_from_pdf(pdf_path):
    if not os.path.exists(pdf_path):
        print("❌ PDF not found!")
        return

    try:
        pages = convert_from_path(pdf_path, poppler_path=POPPLER_PATH)

        for i, page in enumerate(pages):
            text = pytesseract.image_to_string(page)

            aadhar = is_aadhar(text)
            if aadhar:
                name = get_name_near_aadhar(text)
                dob = extract_dob(text)
                gender = extract_gender(text)

                save_to_excel(name, aadhar, dob, gender)

                print(f"\n🆔 Aadhaar detected on page {i+1}:")
                print(f"👤 Name: {name}")
                print(f"📅 DOB: {dob}")
                print(f"🚻 Gender: {gender}")
                print(f"🔢 Aadhaar: {aadhar}")
                return

        print("🚫 No Aadhaar number found!")

    except Exception as e:
        print(f"❌ OCR Error: {e}")

# =======================
# DAY 5 EXTRA FEATURES
# =======================
def search_keyword(file_path, keyword):
    if not os.path.exists(file_path):
        print("❌ File not found!")
        return
    
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        count = text.lower().count(keyword.lower())
        print(f"🔍 Keyword '{keyword}' found {count} times.")
    except Exception as e:
        print(f"❌ Error: {e}")

def analyze_special_characters(file_path):
    if not os.path.exists(file_path):
        print("❌ File not found!")
        return
    
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        special = re.findall(r'[^A-Za-z0-9\s]', text)
        print(f"🎭 Special Characters Found: {len(special)}")
    except Exception as e:
        print(f"❌ Error: {e}")

def export_analysis_to_excel():
    if not last_text_result and not last_csv_result:
        print("⚠ Perform analysis first before export!")
        return

    excel_path = os.path.join(BASE_DIR, "analysis_output.xlsx")
    wb = Workbook()
    ws = wb.active
    ws.title = "Analysis Report"
    ws.append(["Analysis Type", "Details"])

    if last_text_result:
        ws.append(["Text Analysis", last_text_result])
    if last_csv_result is not None:
        ws.append(["CSV Analysis", str(last_csv_result)])

    wb.save(excel_path)
    print(f"📁 Summary exported to Excel at: {excel_path}")

from datetime import datetime

def generate_final_summary(text_summary, csv_summary, ocr_summary="OCR Aadhaar Data extracted (if scanned)"):
    output_path = os.path.join(BASE_DIR, "final_project_summary.txt")

    now = datetime.now().strftime("%d-%m-%Y | Time: %I:%M %p")

    summary = f"""
FILE ANALYZER TOOL – FINAL PROJECT SUMMARY
Developer : Kedar Kumar Trivedi
Date      : {now}
Version   : v1.0 (Mini Project – SEM 4)
-----------------------------------------

TEXT ANALYSIS
{text_summary}

CSV ANALYSIS
{csv_summary}

OCR Aadhaar:
{ocr_summary}

-----------------------------------------
Generated using Python – File Analyzer Tool
    """

    with open(output_path, "w", encoding="utf-8") as file:
        file.write(summary)

    print(f"📄 Final summary saved at: {output_path}")


# =======================
# MAIN MENU
# =======================
def main_menu():
    while True:
        print("\n==== FILE ANALYZER TOOL ====")
        print("1) Analyze Text File & Generate Report")
        print("2) Analyze CSV File")
        print("3) Scan PDF for Aadhaar")
        print("4) Search Keyword in Text File")
        print("5) Analyze Special Characters")
        print("6) Export All Analysis to Excel")
        print("7) Generate Final Summary Report")
        print("8) Exit")

        choice = input("👉 Enter your choice: ")

        if choice == '1':
            analyze_text_file(input("📄 Enter text file path: "))
        elif choice == '2':
            analyze_csv_file(input("📊 Enter CSV file path: "))
        elif choice == '3':
            extract_aadhar_from_pdf(input("📄 Enter PDF file path: "))
        elif choice == '4':
            search_keyword(input("📄 Enter text file path: "), input("🔍 Enter keyword: "))
        elif choice == '5':
            analyze_special_characters(input("📄 Enter text file path: "))
        elif choice == '6':
            export_analysis_to_excel()
        elif choice == '7':
            generate_final_summary()
        elif choice == '8':
            print("🚪 Exiting... Goodbye!")
            break
        else:
            print("⚠ Invalid choice! Try again")

# =======================
# RUN PROGRAM
# =======================
if __name__ == "__main__":
    main_menu()
